from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "/Users/mt/Documents/Codex/outputs/material_retention_word/4月基准对比1月素材侧付费留存分析.docx"


COLORS = {
    "navy": RGBColor(31, 55, 86),
    "blue": RGBColor(46, 91, 140),
    "light_blue": "D9EAF7",
    "pale_blue": "EFF6FB",
    "green": "E2F0D9",
    "red": "FCE4D6",
    "gray": "F2F2F2",
    "dark_gray": RGBColor(89, 89, 89),
    "white": RGBColor(255, 255, 255),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFBFBF", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, pct=5000):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(pct))
    tbl_w.set(qn("w:type"), "pct")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.color.rgb = COLORS["navy"] if level == 1 else COLORS["blue"]
    run.font.bold = True
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    p.add_run(text)
    return p


def add_callout(doc, title, body, fill="EFF6FB"):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    set_table_borders(table, color="D9E2F3", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = COLORS["navy"]
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.15
    p2.add_run(body).font.size = Pt(10)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None, header_fill="D9EAF7"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    set_table_borders(table, color="BFBFBF")
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_shading(c, header_fill)
        set_cell_text(c, h, bold=True, color=COLORS["navy"], size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_margins(c, top=110, bottom=110, start=80, end=80)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]
            set_cell_text(c, value, size=8.8, align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 or i == len(row)-1 else WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_margins(c, top=90, bottom=90, start=80, end=80)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_metric_strip(doc):
    headers = ["指标", "1月", "4月", "变化"]
    rows = [
        ["买量花费", "198,423", "346,184", "+74.5%"],
        ["注册人数", "4,278", "9,311", "+117.7%"],
        ["首日付费人数", "122", "214", "+75.4%"],
        ["首日付费率", "2.9%", "2.3%", "-0.6pp"],
        ["加权首日 ROI", "2.7%", "1.9%", "-0.8pp"],
        ["付费用户 7日留存", "32.3%", "27.5%", "-4.8pp"],
        ["付费用户 14日留存", "23.1%", "17.6%", "-5.5pp"],
    ]
    table = add_table(doc, headers, rows, widths=[4.0, 3.0, 3.0, 3.0], header_fill="D9EAF7")
    for row in table.rows[1:]:
        change = row.cells[3].text
        if change.startswith("+"):
            set_cell_shading(row.cells[3], COLORS["green"])
        elif change.startswith("-"):
            set_cell_shading(row.cells[3], COLORS["red"])


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("4月基准对比1月素材侧付费留存分析")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = COLORS["navy"]

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("分析周期：2026-04-04至2026-04-08 对比 2026-01-13至2026-01-15｜口径：素材级付费用户留存")
    rs.font.size = Pt(10)
    rs.font.color.rgb = COLORS["dark_gray"]

    add_callout(
        doc,
        "总判断",
        "4月不是整体变好，而是买量规模明显放大，但付费留存和首日 ROI 被稀释；少数素材能承接放量，多数新增或放量素材没有形成稳定留存。",
        fill="EFF6FB",
    )

    add_heading(doc, "一、总盘变化", 1)
    add_body(doc, "4月花费、注册和首日付费人数均明显增长，但新增量的付费质量偏弱。买量花费从 19.8 万增加到 34.6 万，注册人数翻倍以上；同时首日付费率、首日 ROI、7日留存、14日留存同步下降。")
    add_metric_strip(doc)

    add_heading(doc, "二、素材分层结论", 1)
    add_heading(doc, "1. 高消耗且可继续承接", 2)
    add_table(
        doc,
        ["素材", "4月花费", "4月 D7", "4月 D14", "判断"],
        [
            ["试玩-护城2代迭代+zc", "47,163", "35.5%", "29.0%", "继续保留，是当前最稳主力"],
            ["试玩-国王养成迭代1+DA+zc", "20,752", "42.9%", "28.6%", "留存较1月明显改善，可继续验证放量"],
            ["试玩-护城2代迭代05+zc", "19,395", "30.0%", "30.0%", "留存合格，但 ROI 低，需要控成本"],
            ["试玩BJ-海岛守卫新传2+XL+zc-251107", "16,669", "37.5%", "25.0%", "放量后没有破坏留存，可小幅加预算"],
            ["试玩-骑士养成迭代新传1+DA+zc", "9,566", "57.1%", "42.9%", "样本小但留存好，适合加预算验证"],
        ],
        widths=[6.4, 2.1, 1.8, 1.8, 4.0],
        header_fill="D9EAF7",
    )

    add_heading(doc, "2. 高消耗但需要降量或重做表达", 2)
    add_table(
        doc,
        ["素材", "4月花费", "ROI", "D7", "D14", "主要问题"],
        [
            ["试玩-护城2代迭代02+zc", "34,037", "1.2%", "24.0%", "8.0%", "1月 D14 为 27.3%，4月明显塌陷"],
            ["ZC-展示+玩法展示+选择战斗3...", "20,981", "5.6%", "20.8%", "12.5%", "首日回收强，但后续留存弱"],
            ["试玩-骑士养成迭代1+DA+zc", "19,183", "0.4%", "0.0%", "0.0%", "从1月 D7 35.3%、D14 29.4% 直接失效"],
            ["XW-Ai+场景展示+国王的家-250724-720", "11,176", "0.5%", "25.0%", "0.0%", "低付费率、低回收，不应继续放大"],
            ["SP-融合玩法+变种塔防+WAW战斗-250523-720", "5,276", "0.0%", "-", "-", "无首日付费，停投优先级高"],
        ],
        widths=[6.2, 1.7, 1.3, 1.3, 1.3, 4.2],
        header_fill="FCE4D6",
    )

    add_heading(doc, "三、共同素材对比", 1)
    add_body(doc, "两期买量共同素材有 68 个，但两期都有留存记录的只有 16 个，因此真正能严肃同比的样本不多。结论应优先看高消耗且 pay 样本足够的素材。")
    add_table(
        doc,
        ["素材", "1月花费", "4月花费", "1月 D7/D14", "4月 D7/D14", "判断"],
        [
            ["试玩-护城2代迭代+zc", "49,433", "47,163", "30.6% / 13.9%", "35.5% / 29.0%", "次留下降，但中后期更稳，是核心素材"],
            ["试玩-国王养成迭代1+DA+zc", "17,009", "20,752", "11.1% / 11.1%", "42.9% / 28.6%", "方向明显变好，但 pay 样本只有 7"],
            ["试玩BJ-海岛守卫新传2+XL+zc-251107", "8,963", "16,669", "20.0% / 20.0%", "37.5% / 25.0%", "放量后留存改善"],
            ["试玩-护城2代迭代02+zc", "32,708", "34,037", "27.3% / 27.3%", "24.0% / 8.0%", "D14 明显塌陷，需要降量排查"],
            ["试玩-骑士养成迭代1+DA+zc", "19,442", "19,183", "35.3% / 29.4%", "0.0% / 0.0%", "明确失效素材"],
        ],
        widths=[5.6, 1.7, 1.7, 2.2, 2.2, 3.6],
        header_fill="D9EAF7",
    )

    add_heading(doc, "四、结构原因", 1)
    add_body(doc, "4月留存下降不是单一素材造成，而是系统、版位和首充金额结构同时改变。新增量集中在更容易拉低后续留存的人群上。")
    add_table(
        doc,
        ["维度", "1月", "4月", "影响判断"],
        [
            ["Android", "pay 76；D7 32.9%；D14 22.4%", "pay 127；D7 25.2%；D14 18.1%", "最大拖累项，新增 Android 量质量偏弱"],
            ["iOS", "pay 50；D7 32.0%；D14 24.0%", "pay 64；D7 34.4%；D14 18.7%", "D7 略升，但 D14 下降"],
            ["激励版位", "pay 17；D7 35.3%；D14 29.4%", "pay 69；D7 15.9%；D14 13.0%", "扩量后留存明显变差"],
            ["激励试玩", "pay 65；D7 26.2%；D14 24.6%", "pay 87；D7 36.8%；D14 19.5%", "7日改善，但14日承接不足"],
            ["低首充 [1,7)", "pay 71；D7 25.4%；D14 19.7%", "pay 140；D7 23.6%；D14 15.7%", "低额付费占比提升，拉低中后期留存"],
        ],
        widths=[3.2, 4.0, 4.0, 5.0],
        header_fill="F2F2F2",
    )

    add_heading(doc, "五、执行建议", 1)
    add_bullet(doc, "保留并继续放量：试玩-护城2代迭代+zc、试玩BJ-海岛守卫新传2+XL+zc-251107、试玩-国王养成迭代1+DA+zc。")
    add_bullet(doc, "小预算验证放量：试玩-骑士养成迭代新传1+DA+zc、way-氛围+红白机甲-251031-1280、试玩KS-护城4代近战+ZC-251128。")
    add_bullet(doc, "降量观察：ZC-展示+玩法展示+选择战斗3...。它首日 ROI 好，但 D7/D14 不够稳，适合短回收，不适合作为长期留存主力。")
    add_bullet(doc, "优先停投或重做表达：试玩-骑士养成迭代1+DA+zc、XW-Ai+场景展示+国王的家-250724-720、SP-融合玩法+变种塔防+WAW战斗-250523-720。")
    add_bullet(doc, "投放策略上，4月的问题不是缺量，而是新增量质量被稀释；后续应把预算从“能买到首日付费”转向“D7/D14 能承接”的素材。")

    add_heading(doc, "六、口径说明", 1)
    add_bullet(doc, "本报告中的留存指付费用户留存，不代表全部注册用户留存。")
    add_bullet(doc, "素材名完全一致时才视为同一素材；不做模糊合并，避免误合并不同迭代素材。")
    add_bullet(doc, "留存表按素材聚合时，使用 pay用户数 加权：sum(pay用户数 × 留存率) / sum(pay用户数)。")
    add_bullet(doc, "买量表中的 '-' 按缺失值处理，不参与对应指标计算。")
    add_bullet(doc, "小样本素材只作为方向参考，不作为强结论。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

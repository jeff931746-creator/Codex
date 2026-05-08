from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path("/Users/mt/Documents/Codex/tmp/level_progress_analysis/data/april_focus")
INPUT = Path("/Users/mt/Downloads/1.13-1.15关卡进度.xlsx")
OUT_DIR = Path("/Users/mt/Documents/Codex/tmp/level_progress_analysis/reports")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "4月关卡流失综合分析文档.docx"

PAY_ORDER = ["1~7", "7~50", "50~100", "100~300", "300~+∞"]
JAN_SHEETS = ["1.13-1.15", "1.09-1.11"]
ALL_SHEETS = ["4.4-4.8", "1.13-1.15", "1.09-1.11"]


def fmt_pct(x, digits=1):
    if pd.isna(x):
        return "-"
    return f"{float(x) * 100:.{digits}f}%"


def fmt_num(x, digits=1):
    if pd.isna(x):
        return "-"
    if abs(float(x) - round(float(x))) < 1e-9:
        return str(int(round(float(x))))
    return f"{float(x):.{digits}f}"


def stage_order(value):
    if value is None or pd.isna(value):
        return float("nan")
    v = int(value)
    if v < 100:
        return float("nan")
    chapter = v // 100
    level = v % 100
    if chapter <= 0 or level <= 0:
        return float("nan")
    return float((chapter - 1) * 15 + level)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color: RGBColor):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color


def set_cell_font(cell, size=9, bold=False):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.bold = bold
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def set_cell_width(cell, width_cm: float):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))


def preferred_widths(col_count: int) -> list[float]:
    if col_count == 1:
        return [25.5]
    if col_count == 2:
        return [5.0, 20.5]
    if col_count == 4:
        return [2.3, 4.0, 13.2, 6.0]
    if col_count == 5:
        return [4.2, 5.3, 5.3, 5.3, 5.4]
    if col_count == 7:
        return [3.2, 3.5, 3.5, 3.5, 3.7, 4.2, 3.9]
    if col_count == 9:
        return [2.2, 2.5, 2.4, 2.4, 2.1, 2.1, 3.1, 2.8, 2.9]
    width = 25.5 / max(col_count, 1)
    return [width] * col_count


def apply_table_widths(table, widths: list[float] | None = None):
    table.autofit = False
    widths = widths or preferred_widths(len(table.columns))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def style_table(table, header_fill="1F4E79", stripe_fill="F6F8FA"):
    table.style = "Table Grid"
    apply_table_widths(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_font(cell, size=8.5 if len(table.columns) > 6 else 9, bold=(row_idx == 0))
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                set_cell_text_color(cell, RGBColor(255, 255, 255))
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif row_idx % 2 == 0:
                set_cell_shading(cell, stripe_fill)


def add_table(doc: Document, title: str, columns: list[str], rows: list[list[str]], note: str | None = None):
    p = doc.add_paragraph()
    p.style = "Heading 3"
    p.add_run(title)
    table = doc.add_table(rows=1, cols=len(columns))
    for i, c in enumerate(columns):
        table.rows[0].cells[i].text = c
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table)
    if note:
        n = doc.add_paragraph(note)
        n.style = "Caption"
    return table


def add_bullets(doc: Document, items: list[str], style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def apply_styles(doc: Document):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.18
    styles["Normal"].paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Title", 22, "17365D"),
        ("Heading 1", 16, "17365D"),
        ("Heading 2", 13, "1F4E79"),
        ("Heading 3", 10.5, "5B6770"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    styles["Caption"].font.name = "Arial"
    styles["Caption"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    styles["Caption"].font.size = Pt(8)
    styles["Caption"].font.color.rgb = RGBColor(90, 90, 90)
    styles["Caption"].paragraph_format.space_after = Pt(6)


def section_metrics(sub: pd.DataFrame):
    n = len(sub)
    churn = int(sub["是否流失"].sum())
    front = int(((sub["最终进度"] >= 1) & (sub["最终进度"] <= 4) & sub["是否流失"]).sum())
    carry = int(((sub["最终进度"] >= 5) & (sub["最终进度"] <= 15) & sub["是否流失"]).sum())
    later = int(((sub["最终进度"] > 15) & sub["是否流失"]).sum())
    return n, churn, front, carry, later


def build_day_distribution(rows: pd.DataFrame):
    churned = rows[rows["是否流失"]].copy()
    bins = [0, 1, 4, 24]
    labels = ["首日", "2-4日", "5日后"]
    churned["流失天数段"] = pd.cut(churned["流失天数"], bins=bins, labels=labels, include_lowest=True, right=True)
    out = []
    for pay in PAY_ORDER:
        for period, mask in [
            ("4月", churned["批次"] == "4.4-4.8"),
            ("1月基线", churned["批次"].isin(JAN_SHEETS)),
        ]:
            sub = churned[mask & (churned["付费分层"] == pay)]
            total = len(sub)
            row = {"付费分层": pay, "时期": period, "流失人数": total, "平均流失天数": sub["流失天数"].mean()}
            for label in labels:
                count = int((sub["流失天数段"] == label).sum())
                row[label] = count
                row[f"{label}占比"] = count / total if total else 0
            out.append(row)
    return pd.DataFrame(out)


def compute_loss_days():
    frames = []
    for sheet in ALL_SHEETS:
        df = pd.read_excel(INPUT, sheet_name=sheet)
        cols = list(df.columns)
        progress_id_idx = cols.index("游戏ID.1")
        raw_cols = cols[2 : progress_id_idx - 2]
        raw = df[raw_cols].apply(pd.to_numeric, errors="coerce").mask(lambda x: x < 100)
        progress = raw.map(stage_order)

        def last_day(series: pd.Series):
            valid_index = series.dropna().index
            if len(valid_index) == 0:
                return float("nan")
            try:
                return int(valid_index[-1]) + 1
            except Exception:
                return float("nan")

        frames.append(
            pd.DataFrame(
                {
                    "批次": sheet,
                    "游戏ID": df[cols[0]],
                    "流失天数": progress.apply(last_day, axis=1),
                }
            )
        )
    return pd.concat(frames, ignore_index=True)


def load_data():
    rows = pd.read_csv(BASE / "player_detail_all.csv")
    loss_days = compute_loss_days()
    rows = rows.merge(loss_days, on=["批次", "游戏ID"], how="left")
    period = pd.read_csv(BASE / "period_summary.csv")
    priority = pd.read_csv(BASE / "priority_levels.csv")
    pay = pd.read_csv(BASE / "pay_comparison.csv")
    return rows, period, priority, pay


def main():
    rows, period, priority, pay = load_data()
    day_dist = build_day_distribution(rows)

    doc = Document()
    apply_styles(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.3)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.3)
    sec.right_margin = Cm(1.3)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("4月关卡流失综合分析文档")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("以 4.4-4.8 为主，对比 1.13-1.15、1.09-1.11；重点关注买量波动、付费分层、流失天数与第一章承接")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(90, 90, 90)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("数据源：1.13-1.15关卡进度.xlsx ｜ 输出日期：2026-04-28")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(110, 110, 110)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    apply_table_widths(callout, [25.5])
    callout_cell = callout.rows[0].cells[0]
    callout_cell.text = (
        "核心结论：4月流失不是单一关卡问题，而是流失窗口整体前移。"
        "入口侧表现为首日早筛变重，主要集中在 1~7 分层；承接侧表现为 2-4日内在第一章后段掉失，"
        "重点关卡是 1-12、1-14，长期卡点是 2-04。"
    )
    set_cell_shading(callout_cell, "EAF3F8")
    set_cell_font(callout_cell, size=10, bold=True)

    doc.add_heading("1. 执行摘要", level=1)
    add_bullets(
        doc,
        [
            "4月整体流失率 68.0%，高于1月基线 59.0%；最终平均进度 19.1，低于1月基线 25.7，说明推进明显变浅。",
            "前四关体验在三期无差别，因此 1-01 到 1-04 只作为流量质量、样本结构和自然早筛的判断口径。",
            "4月新增/加重的关卡承接问题主要在 1-12、1-14、1-11；其中 1-12、1-14 同时具备高流失人数和高归一化恶化。",
            "1~7 是最大损失来源：贡献4月全部流失的约67%，也贡献第一章承接流失的约69%。",
            "按流失天数看，4月流失明显前置：首日占比 44.5%，高于1月基线 35.4%；5日后流失占比反而下降。",
        ],
    )

    doc.add_heading("2. 数据口径与解释约束", level=1)
    add_table(
        doc,
        "表1 口径说明",
        ["项目", "口径"],
        [
            ["主分析对象", "4.4-4.8"],
            ["对比基线", "1.13-1.15、1.09-1.11，并合并为1月基线"],
            ["累计进度", "使用原始关卡号重算：(章节 - 1) * 15 + 小关"],
            ["无效值", "原始关卡号为0的记录按无效处理"],
            ["流失天数", "玩家最后一次有有效关卡记录的天数"],
            ["前四关约束", "1-01 到 1-04 三期体验无差别，只判断买量质量/自然筛选，不做关卡体验归因"],
            ["承接区间", "第一章承接定义为 1-05 到 1-15"],
        ],
    )

    doc.add_heading("3. 整体流失结构", level=1)
    period_rows = []
    for label in ["4.4-4.8", "1月基线", "1.13-1.15", "1.09-1.11"]:
        p = period[period["批次"] == label].iloc[0]
        period_rows.append(
            [
                label,
                int(p["人数"]),
                fmt_pct(p["流失率"]),
                fmt_num(p["最终平均进度"]),
                fmt_num(p["最终中位进度"]),
                fmt_num(p["平均记录天数"]),
                fmt_pct(p["前四关最终停留占比"]),
            ]
        )
    add_table(
        doc,
        "表2 三期整体对比",
        ["批次", "样本", "流失率", "最终均进度", "最终中位", "平均记录天数", "前四关停留占比"],
        period_rows,
        "注：4月最终均进度显著低于1月基线，说明流失和停滞更靠前。",
    )

    split_rows = []
    for label, mask in [
        ("4月", rows["批次"] == "4.4-4.8"),
        ("1月基线", rows["批次"].isin(JAN_SHEETS)),
    ]:
        n, churn, front, carry, later = section_metrics(rows[mask])
        split_rows.append(
            [
                label,
                n,
                churn,
                f"{front} / {fmt_pct(front / n)}",
                f"{carry} / {fmt_pct(carry / n)}",
                f"{later} / {fmt_pct(later / n)}",
            ]
        )
    add_table(
        doc,
        "表3 流失区段拆解",
        ["时期", "样本", "流失人数", "前四关早筛", "第一章承接", "第二章及以后"],
        split_rows,
    )
    add_bullets(
        doc,
        [
            "4月前四关早筛占样本约14.0%，高于1月基线约9.9%，说明入口质量/自然筛选更重。",
            "4月第一章承接占样本约43.0%，高于1月基线约37.4%，说明前四关之后的承接同样恶化。",
        ],
    )

    doc.add_heading("4. 付费分层对比", level=1)
    pay_rows = []
    for tier in PAY_ORDER:
        ap = rows[(rows["批次"] == "4.4-4.8") & (rows["付费分层"] == tier)]
        jan = rows[(rows["批次"].isin(JAN_SHEETS)) & (rows["付费分层"] == tier)]
        a_n, a_churn, a_front, a_carry, a_later = section_metrics(ap)
        j_n, j_churn, j_front, j_carry, j_later = section_metrics(jan)
        pay_rows.append(
            [
                tier,
                a_n,
                fmt_pct(a_churn / a_n if a_n else 0),
                fmt_pct(j_churn / j_n if j_n else 0),
                fmt_pct((a_churn / a_n if a_n else 0) - (j_churn / j_n if j_n else 0)),
                f"{a_front}/{a_carry}/{a_later}",
                fmt_pct(a_carry / a_churn if a_churn else 0),
                fmt_num(ap["最终进度"].mean()),
                fmt_num(jan["最终进度"].mean()),
            ]
        )
    add_table(
        doc,
        "表4 各付费分层流失结构",
        ["分层", "4月样本", "4月流失率", "1月流失率", "变化", "4月早筛/承接/后期", "承接占流失", "4月最终均", "1月最终均"],
        pay_rows,
        "注：承接占流失指 1-05 到 1-15 流失占该分层全部流失的比例。",
    )
    add_bullets(
        doc,
        [
            "1~7 是最大损失来源，4月流失率 75.4%，比1月基线高10.6pct；承接流失占该层流失65.4%。",
            "7~50 的流失率只小幅上升，但2-4日流失集中，说明玩家愿意玩几天后在承接阶段掉失。",
            "100~300 4月首日流失异常前置，但样本只有18人，建议作为渠道/付费触发异常观察，不宜单独下强结论。",
            "50~100 与300+样本过小，只做方向参考。",
        ],
    )

    doc.add_heading("5. 流失天数分析", level=1)
    # Overall day distribution
    day_all_rows = []
    for label, mask in [
        ("4月", rows["批次"] == "4.4-4.8"),
        ("1月基线", rows["批次"].isin(JAN_SHEETS)),
    ]:
        churned = rows[mask & rows["是否流失"]]
        total = len(churned)
        d1 = int((churned["流失天数"] == 1).sum())
        d24 = int(churned["流失天数"].between(2, 4).sum())
        d5 = int((churned["流失天数"] >= 5).sum())
        day_all_rows.append([label, total, f"{d1} / {fmt_pct(d1 / total)}", f"{d24} / {fmt_pct(d24 / total)}", f"{d5} / {fmt_pct(d5 / total)}"])
    add_table(
        doc,
        "表5 整体流失天数分布",
        ["时期", "流失人数", "首日", "2-4日", "5日后"],
        day_all_rows,
    )

    tier_day_rows = []
    for tier in PAY_ORDER:
        for period_label in ["4月", "1月基线"]:
            row = day_dist[(day_dist["付费分层"] == tier) & (day_dist["时期"] == period_label)]
            if row.empty:
                continue
            r = row.iloc[0]
            tier_day_rows.append(
                [
                    tier,
                    period_label,
                    int(r["流失人数"]),
                    f"{int(r['首日'])} / {fmt_pct(r['首日占比'])}",
                    f"{int(r['2-4日'])} / {fmt_pct(r['2-4日占比'])}",
                    f"{int(r['5日后'])} / {fmt_pct(r['5日后占比'])}",
                    fmt_num(r["平均流失天数"]),
                ]
            )
    add_table(
        doc,
        "表6 分层流失天数分布",
        ["分层", "时期", "流失人数", "首日", "2-4日", "5日后", "平均流失天数"],
        tier_day_rows,
        "注：1~7贡献4月首日流失的约74%；7~50的4月流失主要集中在2-4日。",
    )
    add_bullets(
        doc,
        [
            "4月流失窗口整体前移：首日流失占44.5%，比1月基线高9.1pct。",
            "1~7 是首日早筛主因：4月首日流失51人，占该层流失49.0%。",
            "7~50 主要在2-4日掉失：4月2-4日流失19人，占该层流失52.8%。",
            "5日后不是4月主要矛盾，占比低于1月基线。",
        ],
    )

    doc.add_heading("6. 关键关卡与影响判断", level=1)
    top = priority[~priority["是否前四关基准段"]].head(10)
    level_rows = []
    for _, r in top.iterrows():
        level_rows.append(
            [
                r["关卡标签"],
                r["问题类型"],
                int(r["4月到达人数"]),
                int(r["4月流失人数"]),
                fmt_pct(r["4月归一化流失率"]),
                fmt_pct(r["1月基线归一化流失率"]),
                fmt_pct(r["4月变化"]),
                fmt_num(r["优先级分"]),
            ]
        )
    add_table(
        doc,
        "表7 4月关键卡点优先级",
        ["关卡", "类型", "4月到达", "4月流失", "4月流失率", "1月基线", "变化", "优先级"],
        level_rows,
    )
    add_bullets(
        doc,
        [
            "1-12、1-14是最明确的4月新增/加重卡点，既有绝对人数，也有归一化恶化。",
            "2-04 是长期卡点，4月没有大幅恶化但仍有较高归一化流失率。",
            "1-03、1-04虽有较高早期流失，但按约束只作为流量质量/自然筛选信号。",
        ],
    )

    doc.add_heading("7. 行动建议", level=1)
    add_table(
        doc,
        "表8 建议优先级",
        ["优先级", "方向", "动作", "验证指标"],
        [
            ["P0", "买量质量与首日预期", "按渠道/素材/计划拆首日流失，重点看1~7首日掉失；检查素材承诺与实际玩法是否错配。", "首日流失占比、1-03/1-04停留率、首日有效记录率"],
            ["P0", "第一章后段承接", "复盘1-12、1-14的怪物强度、资源供给、技能/装备解锁、失败后补强路径。", "1-12/1-14归一化流失率、2-4日流失占比"],
            ["P1", "7~50承接", "针对7~50做2-4日关卡路径检查，确认二章初期是否出现战力断档。", "7~50的2-4日流失、2-04流失率"],
            ["P1", "100~300异常样本", "回看4月高付费早流失用户来源和充值触发点，判断是否为小样本渠道异常。", "100~300首日流失、来源集中度"],
            ["P2", "长期卡点", "持续观察2-04、1-09、1-10、1-13，避免只修新增点导致老卡点继续吞量。", "长期卡点归一化流失率"],
        ],
    )

    doc.add_heading("8. 结论", level=1)
    add_bullets(
        doc,
        [
            "4月的流失恶化由两部分叠加：入口质量/早筛变重，以及第一章后段承接变差。",
            "从人数和影响看，1~7 是必须优先处理的核心分层；从流失天数看，首日与2-4日是关键窗口。",
            "从关卡看，1-12、1-14是最值得立即复盘的新增/加重问题；2-04属于长期卡点，需持续优化。",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

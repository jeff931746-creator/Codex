from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path("/Users/mt/Documents/Codex/tmp/level_progress_analysis/data/april_focus")
INPUT = Path("/Users/mt/Downloads/1.13-1.15关卡进度.xlsx")
OUT_DIR = Path("/Users/mt/Documents/Codex/tmp/level_progress_analysis/reports")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "4月关卡流失综合分析文档.docx"

PAY_ORDER = ["1~7", "7~50", "50~100", "100~300", "300~+∞"]
JAN_SHEETS = ["1.13-1.15", "1.09-1.11"]
ALL_SHEETS = ["4.4-4.8", "1.13-1.15", "1.09-1.11"]


def pct(x, digits=1):
    if pd.isna(x):
        return "-"
    return f"{float(x) * 100:.{digits}f}%"


def num(x, digits=1):
    if pd.isna(x):
        return "-"
    if abs(float(x) - round(float(x))) < 1e-9:
        return str(int(round(float(x))))
    return f"{float(x):.{digits}f}"


def stage_order(value):
    if value is None or pd.isna(value):
        return float("nan")
    v = int(value)
    if v < 100:
        return float("nan")
    chapter = v // 100
    level = v % 100
    if chapter <= 0 or level <= 0:
        return float("nan")
    return float((chapter - 1) * 15 + level)


def compute_loss_days():
    frames = []
    for sheet in ALL_SHEETS:
        df = pd.read_excel(INPUT, sheet_name=sheet)
        cols = list(df.columns)
        progress_id_idx = cols.index("游戏ID.1")
        raw_cols = cols[2 : progress_id_idx - 2]
        raw = df[raw_cols].apply(pd.to_numeric, errors="coerce").mask(lambda x: x < 100)
        progress = raw.map(stage_order)

        def last_day(series):
            valid_index = series.dropna().index
            if len(valid_index) == 0:
                return float("nan")
            return int(valid_index[-1]) + 1

        frames.append(pd.DataFrame({"批次": sheet, "游戏ID": df[cols[0]], "流失天数": progress.apply(last_day, axis=1)}))
    return pd.concat(frames, ignore_index=True)


def para_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def apply_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.18
    styles["Normal"].paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Title", 22, "17365D"),
        ("Heading 1", 16, "17365D"),
        ("Heading 2", 13, "1F4E79"),
        ("Heading 3", 11, "5B6770"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def add_label_line(doc, label, text, shade=None):
    p = doc.add_paragraph()
    if shade:
        para_shading(p, shade)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def section_metrics(sub):
    n = len(sub)
    churn = int(sub["是否流失"].sum())
    front = int(((sub["最终进度"] >= 1) & (sub["最终进度"] <= 4) & sub["是否流失"]).sum())
    carry = int(((sub["最终进度"] >= 5) & (sub["最终进度"] <= 15) & sub["是否流失"]).sum())
    later = int(((sub["最终进度"] > 15) & sub["是否流失"]).sum())
    return n, churn, front, carry, later


def load_data():
    rows = pd.read_csv(BASE / "player_detail_all.csv")
    rows = rows.merge(compute_loss_days(), on=["批次", "游戏ID"], how="left")
    period = pd.read_csv(BASE / "period_summary.csv")
    priority = pd.read_csv(BASE / "priority_levels.csv")
    return rows, period, priority


def main():
    rows, period, priority = load_data()
    doc = Document()
    apply_styles(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.4)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("4月关卡流失综合分析文档")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("以 4.4-4.8 为主，对比 1.13-1.15、1.09-1.11；重点关注买量波动、付费分层、流失天数与第一章承接")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(90, 90, 90)
    meta = doc.add_paragraph("数据源：1.13-1.15关卡进度.xlsx ｜ 输出日期：2026-04-28")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    callout = doc.add_paragraph()
    para_shading(callout, "EAF3F8")
    rr = callout.add_run("核心结论：")
    rr.bold = True
    rr.font.color.rgb = RGBColor(31, 78, 121)
    callout.add_run("4月流失不是单一关卡问题，而是流失窗口整体前移。入口侧表现为首日早筛变重，主要集中在 1~7 分层；承接侧表现为 2-4日内在第一章后段掉失，重点关卡是 1-12、1-14，长期卡点是 2-04。")

    doc.add_heading("1. 执行摘要", level=1)
    add_bullets(
        doc,
        [
            "4月整体流失率 68.0%，高于1月基线 59.0%；最终平均进度 19.1，低于1月基线 25.7。",
            "前四关体验在三期无差别，因此 1-01 到 1-04 只作为流量质量、样本结构和自然早筛的判断口径。",
            "4月新增/加重的关卡承接问题主要在 1-12、1-14、1-11；其中 1-12、1-14 同时具备高流失人数和高归一化恶化。",
            "1~7 是最大损失来源：贡献4月全部流失的约67%，也贡献第一章承接流失的约69%。",
            "按流失天数看，4月流失明显前置：首日占比 44.5%，高于1月基线 35.4%；5日及以后流失占比下降。",
        ],
    )

    doc.add_heading("2. 数据口径与解释约束", level=1)
    for label, text in [
        ("主分析对象：", "4.4-4.8。"),
        ("对比基线：", "1.13-1.15、1.09-1.11，并合并为1月基线。"),
        ("累计进度：", "使用原始关卡号重算：(章节 - 1) * 15 + 小关；原表中间累计进度列不作为主分析依据。"),
        ("流失天数：", "玩家最后一次有有效关卡记录的天数。"),
        ("前四关约束：", "1-01 到 1-04 三期体验无差别，只判断买量质量/自然筛选，不做关卡体验归因。"),
        ("承接区间：", "第一章承接定义为 1-05 到 1-15。"),
    ]:
        add_label_line(doc, label, text)

    doc.add_heading("3. 整体流失结构", level=1)
    for label in ["4.4-4.8", "1月基线", "1.13-1.15", "1.09-1.11"]:
        p = period[period["批次"] == label].iloc[0]
        add_label_line(
            doc,
            f"{label}：",
            f"样本 {int(p['人数'])}；流失率 {pct(p['流失率'])}；最终平均进度 {num(p['最终平均进度'])}；最终中位进度 {num(p['最终中位进度'])}；平均记录天数 {num(p['平均记录天数'])}；前四关最终停留占比 {pct(p['前四关最终停留占比'])}。",
        )
    for label, mask in [("4月", rows["批次"] == "4.4-4.8"), ("1月基线", rows["批次"].isin(JAN_SHEETS))]:
        n, churn, front, carry, later = section_metrics(rows[mask])
        add_label_line(
            doc,
            f"{label}区段拆解：",
            f"样本 {n}，流失 {churn}；前四关早筛 {front}（{pct(front/n)}），第一章承接 {carry}（{pct(carry/n)}），第二章及以后 {later}（{pct(later/n)}）。",
            shade="F6F8FA",
        )

    doc.add_heading("4. 付费分层对比", level=1)
    for tier in PAY_ORDER:
        ap = rows[(rows["批次"] == "4.4-4.8") & (rows["付费分层"] == tier)]
        jan = rows[(rows["批次"].isin(JAN_SHEETS)) & (rows["付费分层"] == tier)]
        a_n, a_churn, a_front, a_carry, a_later = section_metrics(ap)
        j_n, j_churn, _, _, _ = section_metrics(jan)
        if a_n == 0:
            continue
        add_label_line(
            doc,
            f"{tier}：",
            f"4月样本 {a_n}，流失率 {pct(a_churn/a_n if a_n else 0)}，1月流失率 {pct(j_churn/j_n if j_n else 0)}；4月流失结构为前四关 {a_front}、第一章承接 {a_carry}、二章后 {a_later}；最终平均进度 {num(ap['最终进度'].mean())}，1月为 {num(jan['最终进度'].mean())}。",
        )
    add_bullets(
        doc,
        [
            "1~7 是必须优先处理的核心分层，既是人数最大来源，也是第一章承接损失最大的来源。",
            "7~50 不是首日流失主导，而是2-4日承接问题更明显。",
            "100~300 有首日异常前置迹象，但样本偏小，建议作为渠道/付费触发异常观察。",
            "50~100 与300+样本较小，不做强判断。",
        ],
    )

    doc.add_heading("5. 流失天数分析", level=1)
    for label, mask in [("4月", rows["批次"] == "4.4-4.8"), ("1月基线", rows["批次"].isin(JAN_SHEETS))]:
        churned = rows[mask & (rows["是否流失"] == 1)]
        total = len(churned)
        d1 = int((churned["流失天数"] == 1).sum())
        d24 = int(churned["流失天数"].between(2, 4).sum())
        d5 = int((churned["流失天数"] >= 5).sum())
        unknown = int(churned["流失天数"].isna().sum())
        unknown_text = f"，未记录有效关卡 {unknown}（{pct(unknown/total)}）" if unknown else ""
        add_label_line(doc, f"{label}：", f"流失 {total}；首日 {d1}（{pct(d1/total)}），2-4日 {d24}（{pct(d24/total)}），5日及以后 {d5}（{pct(d5/total)}）{unknown_text}。")

    add_label_line(doc, "4月分层贡献：", "1~7贡献首日流失51人，7~50贡献2-4日流失19人，100~300首日流失7人但样本较小。", shade="FFF4E5")
    add_bullets(
        doc,
        [
            "4月流失窗口整体前移：首日流失占44.5%，比1月基线高9.1pct。",
            "1~7 是首日早筛主因：4月首日流失51人，占该层流失49.0%。",
            "7~50 主要在2-4日掉失：4月2-4日流失19人，占该层流失52.8%。",
            "5日及以后不是4月主要矛盾，核心损失已经在前4天发生。",
        ],
    )

    doc.add_heading("6. 关键关卡与影响判断", level=1)
    key = priority[~priority["是否前四关基准段"]].head(10)
    for _, r in key.iterrows():
        add_label_line(
            doc,
            f"{r['关卡标签']}：",
            f"{r['问题类型']}；4月到达 {int(r['4月到达人数'])}，4月流失 {int(r['4月流失人数'])}，4月归一化流失率 {pct(r['4月归一化流失率'])}，1月基线 {pct(r['1月基线归一化流失率'])}，变化 {pct(r['4月变化'])}。",
        )
    add_bullets(
        doc,
        [
            "1-12、1-14 是最明确的4月新增/加重卡点。",
            "2-04 是长期卡点，4月没有大幅恶化但仍有较高归一化流失率。",
            "1-03、1-04 虽有较高早期流失，但按约束只作为买量质量/自然筛选信号。",
        ],
    )

    doc.add_heading("7. 行动建议", level=1)
    for label, text in [
        ("P0 买量质量与首日预期：", "按渠道、素材、计划拆首日流失，重点看1~7首日掉失；检查素材承诺与实际玩法是否错配。"),
        ("P0 第一章后段承接：", "复盘1-12、1-14的怪物强度、资源供给、技能/装备解锁、失败后补强路径。"),
        ("P1 7~50承接：", "针对7~50做2-4日路径检查，确认第一章后段到第二章初期是否出现战力断档。"),
        ("P1 高付费小样本异常：", "回看100~300的4月高付费早流失用户来源和充值触发点，判断是否为渠道异常。"),
        ("P2 长期卡点：", "持续观察2-04、1-09、1-10、1-13，避免只修新增点导致老卡点继续吞量。"),
    ]:
        add_label_line(doc, label, text)

    doc.add_heading("8. 最终结论", level=1)
    add_bullets(
        doc,
        [
            "4月的流失恶化由两部分叠加：入口质量/早筛变重，以及第一章后段承接变差。",
            "从人数和影响看，1~7 是必须优先处理的核心分层；从流失天数看，首日与2-4日是关键窗口。",
            "从关卡看，1-12、1-14 是最值得立即复盘的新增/加重问题；2-04 属于长期卡点，需持续优化。",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
